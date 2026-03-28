import io
import unittest
from unittest.mock import patch

from fastapi import HTTPException
from docx import Document
from pypdf import PdfWriter

from main import (
    ParseRequest,
    decode_text_bytes,
    parse_docx_document,
    parse_document,
    parse_pdf_document,
    parse_text_document,
)


class ParserMainTestCase(unittest.TestCase):
    def test_decode_text_bytes_supports_gb18030(self):
        payload = "合同范围".encode("gb18030")
        self.assertEqual(decode_text_bytes(payload), "合同范围")

    def test_parse_text_document_recognizes_markdown_heading(self):
        result = parse_text_document(b"# \xe5\x90\x88\xe5\x90\x8c\xe8\x8c\x83\xe5\x9b\xb4\n\n\xe7\xac\xac8\xe6\x9d\xa1 \xe4\xb8\x8d\xe5\x8f\xaf\xe6\x8a\x97\xe5\x8a\x9b")

        self.assertEqual(result["page_count"], 1)
        self.assertEqual(result["blocks"][0]["block_type"], "heading")
        self.assertEqual(result["blocks"][0]["text"], "合同范围")

    def test_parse_docx_document_preserves_heading_styles(self):
        document = Document()
        document.add_heading("风险提示", level=1)
        document.add_heading("适用范围", level=2)
        document.add_paragraph("发生不可抗力时应及时通知。")
        buffer = io.BytesIO()
        document.save(buffer)

        result = parse_docx_document(buffer.getvalue())

        self.assertEqual(result["blocks"][0]["block_type"], "heading")
        self.assertEqual(result["blocks"][0]["text"], "风险提示")
        self.assertEqual(result["blocks"][1]["heading_path"], ["风险提示", "适用范围"])
        self.assertEqual(result["blocks"][2]["heading_path"], ["风险提示", "适用范围"])

    def test_parse_docx_document_extracts_tables_in_order(self):
        document = Document()
        document.add_heading("付款计划", level=1)
        table = document.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "阶段"
        table.cell(0, 1).text = "金额"
        table.cell(1, 0).text = "首付款"
        table.cell(1, 1).text = "100万"
        document.add_paragraph("付款应在发票开具后五日内完成。")
        buffer = io.BytesIO()
        document.save(buffer)

        result = parse_docx_document(buffer.getvalue())

        self.assertEqual(result["blocks"][1]["block_type"], "table")
        self.assertIn("阶段 | 金额", result["blocks"][1]["text"])
        self.assertEqual(result["blocks"][1]["heading_path"], ["付款计划"])
        self.assertEqual(result["blocks"][2]["heading_path"], ["付款计划"])

    def test_parse_pdf_document_rejects_blank_text_pdf_without_ocr(self):
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=300)
        buffer = io.BytesIO()
        writer.write(buffer)

        with self.assertRaises(HTTPException) as context:
            parse_pdf_document(buffer.getvalue())

        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(context.exception.detail["code"], "ocr_required")
        self.assertEqual(context.exception.detail["ocr_provider"], "disabled")

    def test_parse_pdf_document_uses_mock_ocr_provider_for_blank_pdf(self):
        writer = PdfWriter()
        writer.add_blank_page(width=300, height=300)
        buffer = io.BytesIO()
        writer.write(buffer)

        with patch.dict(
            "os.environ",
            {
                "PARSER_OCR_PROVIDER": "mock",
                "PARSER_OCR_MOCK_TEXT": "# 合同概述\n\n第8条 不可抗力\n\n发生不可抗力时应及时通知。",
            },
            clear=False,
        ):
            result = parse_pdf_document(buffer.getvalue())

        self.assertEqual(result["source"]["mode"], "ocr")
        self.assertEqual(result["source"]["ocr_provider"], "mock")
        self.assertEqual(result["blocks"][0]["block_type"], "heading")
        self.assertEqual(result["blocks"][1]["heading_path"], ["合同概述", "第8条 不可抗力"])
        self.assertEqual(result["blocks"][2]["heading_path"], ["合同概述", "第8条 不可抗力"])

    def test_parse_document_reads_storage_and_dispatches_by_logical_path(self):
        request = ParseRequest(
            workspace_id="ws_123",
            document_version_id="dv_123",
            storage_key="workspaces/ws_123/uploads/source.bin",
            sha256="abc123",
            logical_path="法规库/notes.md",
        )

        with patch("main.get_object_bytes", return_value="# 合同范围".encode("utf-8")) as get_object_bytes:
            with patch("main.parse_text_document", return_value={"ok": True}) as parse_text:
                result = parse_document(request)

        get_object_bytes.assert_called_once_with("workspaces/ws_123/uploads/source.bin")
        parse_text.assert_called_once_with("# 合同范围".encode("utf-8"))
        self.assertEqual(result, {"ok": True})


if __name__ == "__main__":
    unittest.main()
